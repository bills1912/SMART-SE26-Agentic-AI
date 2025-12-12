from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
    PageBreak, ListFlowable, ListItem, KeepTogether
)
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import logging
import json

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generate PDF and Word reports from chat session analysis data with visualizations and policies"""
    
    def __init__(self):
        self.primary_color = HexColor('#e74c3c')
        self.secondary_color = HexColor('#3498db')
        self.success_color = HexColor('#2ecc71')
        self.warning_color = HexColor('#f39c12')
        self.text_color = HexColor('#2c3e50')
        self.light_bg = HexColor('#f8f9fa')
    
    def _extract_session_data(self, session) -> dict:
        """Extract all relevant data from session for report"""
        data = {
            'title': getattr(session, 'title', 'Analisis Sensus Ekonomi'),
            'messages': [],
            'visualizations': [],
            'insights': [],
            'policies': []
        }
        
        messages = getattr(session, 'messages', [])
        for msg in messages:
            msg_data = {
                'sender': getattr(msg, 'sender', 'user'),
                'content': getattr(msg, 'content', ''),
                'timestamp': getattr(msg, 'timestamp', None)
            }
            
            # Extract visualizations from AI messages
            if msg_data['sender'] == 'ai':
                viz_list = getattr(msg, 'visualizations', [])
                if viz_list:
                    for viz in viz_list:
                        if isinstance(viz, dict):
                            data['visualizations'].append(viz)
                        elif hasattr(viz, 'dict'):
                            data['visualizations'].append(viz.dict())
                        elif hasattr(viz, 'title'):
                            data['visualizations'].append({
                                'title': getattr(viz, 'title', ''),
                                'type': getattr(viz, 'type', 'chart'),
                                'config': getattr(viz, 'config', {}),
                                'data': getattr(viz, 'data', {})
                            })
                
                # Extract insights
                insights_list = getattr(msg, 'insights', [])
                if insights_list:
                    for insight in insights_list:
                        if isinstance(insight, str):
                            data['insights'].append(insight)
                        elif isinstance(insight, dict):
                            data['insights'].append(insight.get('text', insight.get('description', str(insight))))
                
                # Extract policies
                policies_list = getattr(msg, 'policies', [])
                if policies_list:
                    for policy in policies_list:
                        if isinstance(policy, dict):
                            data['policies'].append(policy)
                        elif hasattr(policy, 'dict'):
                            data['policies'].append(policy.dict())
                        elif hasattr(policy, 'title'):
                            data['policies'].append({
                                'title': getattr(policy, 'title', ''),
                                'description': getattr(policy, 'description', ''),
                                'priority': getattr(policy, 'priority', 'medium'),
                                'category': str(getattr(policy, 'category', 'economic')),
                                'impact': getattr(policy, 'impact', ''),
                                'implementation_steps': getattr(policy, 'implementation_steps', [])
                            })
            
            data['messages'].append(msg_data)
        
        return data
    
    def _extract_chart_data_summary(self, viz: dict) -> list:
        """Extract data summary from visualization config for table display"""
        config = viz.get('config', {})
        data_summary = []
        
        try:
            # Try to get data from series
            series = config.get('series', [])
            if series:
                series_data = series[0] if isinstance(series, list) else series
                chart_data = series_data.get('data', [])
                
                # Get categories from xAxis or yAxis
                x_axis = config.get('xAxis', {})
                y_axis = config.get('yAxis', {})
                
                categories = []
                if isinstance(x_axis, dict):
                    categories = x_axis.get('data', [])
                elif isinstance(y_axis, dict):
                    categories = y_axis.get('data', [])
                
                # Build data summary
                if categories and chart_data:
                    for i, cat in enumerate(categories[:10]):  # Limit to 10 rows
                        if i < len(chart_data):
                            value = chart_data[i]
                            if isinstance(value, dict):
                                value = value.get('value', 0)
                            data_summary.append({
                                'label': str(cat),
                                'value': value
                            })
                elif chart_data:
                    # Pie chart data format
                    for item in chart_data[:10]:
                        if isinstance(item, dict):
                            data_summary.append({
                                'label': item.get('name', ''),
                                'value': item.get('value', 0)
                            })
        except Exception as e:
            logger.error(f"Error extracting chart data: {e}")
        
        return data_summary
    
    def generate_pdf(self, session) -> BytesIO:
        """Generate comprehensive PDF report with visualizations and policies"""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4,
                rightMargin=50, 
                leftMargin=50,
                topMargin=50, 
                bottomMargin=50
            )
            
            elements = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            styles.add(ParagraphStyle(
                name='CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=20,
                spaceBefore=10,
                alignment=TA_CENTER,
                textColor=self.primary_color
            ))
            
            styles.add(ParagraphStyle(
                name='SectionTitle',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=10,
                spaceBefore=15,
                textColor=self.primary_color,
                borderPadding=(5, 5, 5, 5)
            ))
            
            styles.add(ParagraphStyle(
                name='SubSection',
                parent=styles['Heading3'],
                fontSize=12,
                spaceAfter=8,
                spaceBefore=10,
                textColor=self.secondary_color
            ))
            
            styles.add(ParagraphStyle(
                name='BodyTextCustom',
                parent=styles['BodyText'],
                fontSize=10,
                spaceAfter=8,
                alignment=TA_JUSTIFY,
                leading=14
            ))
            
            styles.add(ParagraphStyle(
                name='InsightText',
                parent=styles['BodyText'],
                fontSize=10,
                spaceAfter=6,
                leftIndent=20,
                bulletIndent=10
            ))
            
            styles.add(ParagraphStyle(
                name='PolicyTitle',
                parent=styles['Heading4'],
                fontSize=11,
                spaceAfter=4,
                spaceBefore=8,
                textColor=self.text_color
            ))
            
            # Extract session data
            session_data = self._extract_session_data(session)
            
            # === TITLE ===
            title = Paragraph("📊 Laporan Analisis Sensus Ekonomi Indonesia", styles['CustomTitle'])
            elements.append(title)
            elements.append(Spacer(1, 0.2*inch))
            
            # === METADATA ===
            date_str = datetime.now().strftime("%d %B %Y, %H:%M WIB")
            meta_data = [
                ['Tanggal Pembuatan:', date_str],
                ['Topik Analisis:', session_data['title']],
                ['Jumlah Visualisasi:', str(len(session_data['visualizations']))],
                ['Jumlah Rekomendasi:', str(len(session_data['policies']))]
            ]
            
            meta_table = Table(meta_data, colWidths=[150, 300])
            meta_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('TEXTCOLOR', (0, 0), (0, -1), self.text_color),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ]))
            elements.append(meta_table)
            elements.append(Spacer(1, 0.3*inch))
            
            # === RINGKASAN PERCAKAPAN ===
            elements.append(Paragraph("📝 Ringkasan Analisis", styles['SectionTitle']))
            elements.append(Spacer(1, 0.1*inch))
            
            for msg in session_data['messages']:
                sender = "👤 Pengguna" if msg['sender'] == 'user' else "🤖 AI Asisten"
                elements.append(Paragraph(f"<b>{sender}:</b>", styles['SubSection']))
                
                # Clean and format content
                content = msg['content'].replace('\n', '<br/>')
                content = content.replace('**', '<b>').replace('**', '</b>')
                elements.append(Paragraph(content, styles['BodyTextCustom']))
                elements.append(Spacer(1, 0.15*inch))
            
            # === VISUALISASI DATA ===
            if session_data['visualizations']:
                elements.append(PageBreak())
                elements.append(Paragraph("📈 Data Visualisasi", styles['SectionTitle']))
                elements.append(Spacer(1, 0.1*inch))
                
                elements.append(Paragraph(
                    "<i>Catatan: Grafik interaktif tersedia di aplikasi web. "
                    "Berikut adalah ringkasan data dari setiap visualisasi:</i>",
                    styles['BodyTextCustom']
                ))
                elements.append(Spacer(1, 0.15*inch))
                
                for i, viz in enumerate(session_data['visualizations'], 1):
                    viz_title = viz.get('title', f'Visualisasi {i}')
                    elements.append(Paragraph(f"📊 {i}. {viz_title}", styles['SubSection']))
                    
                    # Extract and display data summary
                    data_summary = self._extract_chart_data_summary(viz)
                    
                    if data_summary:
                        # Create data table
                        table_data = [['No.', 'Kategori', 'Jumlah Usaha']]
                        total = 0
                        for j, item in enumerate(data_summary, 1):
                            value = item['value']
                            if isinstance(value, (int, float)):
                                total += value
                                table_data.append([str(j), item['label'], f"{int(value):,}"])
                            else:
                                table_data.append([str(j), item['label'], str(value)])
                        
                        # Add total row if numeric
                        if total > 0:
                            table_data.append(['', 'TOTAL', f"{int(total):,}"])
                        
                        viz_table = Table(table_data, colWidths=[40, 250, 120])
                        viz_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), self.primary_color),
                            ('TEXTCOLOR', (0, 0), (-1, 0), white),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('ALIGN', (2, 0), (2, -1), 'RIGHT'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                            ('ROWBACKGROUNDS', (0, 1), (-1, -2), [white, self.light_bg]),
                            ('BACKGROUND', (0, -1), (-1, -1), HexColor('#ecf0f1')),
                            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ]))
                        elements.append(viz_table)
                    else:
                        elements.append(Paragraph(
                            f"<i>Tipe: {viz.get('type', 'chart')}</i>",
                            styles['BodyTextCustom']
                        ))
                    
                    elements.append(Spacer(1, 0.2*inch))
            
            # === INSIGHTS ===
            if session_data['insights']:
                elements.append(PageBreak())
                elements.append(Paragraph("💡 Insight Analisis", styles['SectionTitle']))
                elements.append(Spacer(1, 0.1*inch))
                
                for i, insight in enumerate(session_data['insights'], 1):
                    insight_text = insight if isinstance(insight, str) else str(insight)
                    elements.append(Paragraph(
                        f"<b>{i}.</b> {insight_text}",
                        styles['InsightText']
                    ))
                    elements.append(Spacer(1, 0.05*inch))
            
            # === POLICY RECOMMENDATIONS ===
            if session_data['policies']:
                elements.append(PageBreak())
                elements.append(Paragraph("🎯 Rekomendasi Kebijakan", styles['SectionTitle']))
                elements.append(Spacer(1, 0.1*inch))
                
                for i, policy in enumerate(session_data['policies'], 1):
                    # Policy title with priority badge
                    priority = policy.get('priority', 'medium')
                    priority_colors = {
                        'high': '🔴',
                        'medium': '🟡', 
                        'low': '🟢'
                    }
                    priority_badge = priority_colors.get(priority, '🟡')
                    
                    policy_title = f"{priority_badge} {i}. {policy.get('title', 'Rekomendasi')}"
                    elements.append(Paragraph(policy_title, styles['PolicyTitle']))
                    
                    # Description
                    description = policy.get('description', '')
                    if description:
                        elements.append(Paragraph(description, styles['BodyTextCustom']))
                    
                    # Category and Impact
                    category = policy.get('category', '')
                    impact = policy.get('impact', '')
                    
                    if category or impact:
                        meta_text = []
                        if category:
                            meta_text.append(f"<b>Kategori:</b> {category}")
                        if impact:
                            meta_text.append(f"<b>Dampak:</b> {impact}")
                        elements.append(Paragraph(' | '.join(meta_text), styles['BodyTextCustom']))
                    
                    # Implementation steps
                    steps = policy.get('implementation_steps', [])
                    if steps:
                        elements.append(Paragraph("<b>Langkah Implementasi:</b>", styles['BodyTextCustom']))
                        for j, step in enumerate(steps, 1):
                            elements.append(Paragraph(
                                f"    {j}. {step}",
                                styles['InsightText']
                            ))
                    
                    elements.append(Spacer(1, 0.15*inch))
            
            # === FOOTER ===
            elements.append(Spacer(1, 0.3*inch))
            elements.append(Paragraph(
                "─" * 60,
                styles['BodyTextCustom']
            ))
            elements.append(Paragraph(
                f"<i>Laporan ini dihasilkan secara otomatis oleh Sistem Analisis Sensus Ekonomi Indonesia.<br/>"
                f"Data bersumber dari Sensus Ekonomi 2016, Badan Pusat Statistik (BPS).<br/>"
                f"© {datetime.now().year} - Smart SE26 Agentic AI Chatbot</i>",
                ParagraphStyle(
                    'Footer',
                    parent=styles['BodyText'],
                    fontSize=8,
                    alignment=TA_CENTER,
                    textColor=colors.grey
                )
            ))
            
            doc.build(elements)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}", exc_info=True)
            raise
    
    def generate_docx(self, session) -> BytesIO:
        """Generate comprehensive Word document with visualizations and policies"""
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
            
            # Extract session data
            session_data = self._extract_session_data(session)
            
            # === TITLE ===
            title = doc.add_heading('Laporan Analisis Sensus Ekonomi Indonesia', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title.runs:
                run.font.color.rgb = RGBColor(231, 76, 60)
            
            # === METADATA ===
            doc.add_paragraph()
            date_str = datetime.now().strftime("%d %B %Y, %H:%M WIB")
            
            meta_table = doc.add_table(rows=4, cols=2)
            meta_table.style = 'Table Grid'
            
            meta_data = [
                ('Tanggal Pembuatan', date_str),
                ('Topik Analisis', session_data['title']),
                ('Jumlah Visualisasi', str(len(session_data['visualizations']))),
                ('Jumlah Rekomendasi', str(len(session_data['policies'])))
            ]
            
            for i, (label, value) in enumerate(meta_data):
                row = meta_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                row.cells[0].paragraphs[0].runs[0].bold = True
            
            doc.add_paragraph()
            
            # === RINGKASAN ANALISIS ===
            doc.add_heading('📝 Ringkasan Analisis', level=1)
            
            for msg in session_data['messages']:
                sender = "👤 Pengguna" if msg['sender'] == 'user' else "🤖 AI Asisten"
                
                p = doc.add_paragraph()
                runner = p.add_run(f"{sender}:")
                runner.bold = True
                runner.font.size = Pt(11)
                
                content_p = doc.add_paragraph(msg['content'])
                content_p.paragraph_format.left_indent = Inches(0.25)
                doc.add_paragraph()
            
            # === VISUALISASI DATA ===
            if session_data['visualizations']:
                doc.add_page_break()
                doc.add_heading('📈 Data Visualisasi', level=1)
                
                note = doc.add_paragraph()
                note_run = note.add_run(
                    'Catatan: Grafik interaktif tersedia di aplikasi web. '
                    'Berikut adalah ringkasan data dari setiap visualisasi:'
                )
                note_run.italic = True
                note_run.font.size = Pt(10)
                
                for i, viz in enumerate(session_data['visualizations'], 1):
                    viz_title = viz.get('title', f'Visualisasi {i}')
                    doc.add_heading(f'📊 {i}. {viz_title}', level=2)
                    
                    # Extract and display data summary
                    data_summary = self._extract_chart_data_summary(viz)
                    
                    if data_summary:
                        # Create data table
                        table = doc.add_table(rows=len(data_summary) + 2, cols=3)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        # Header row
                        header_cells = table.rows[0].cells
                        header_cells[0].text = 'No.'
                        header_cells[1].text = 'Kategori'
                        header_cells[2].text = 'Jumlah Usaha'
                        
                        for cell in header_cells:
                            cell.paragraphs[0].runs[0].bold = True
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), 'E74C3C')
                            cell._tc.get_or_add_tcPr().append(shading)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        
                        # Data rows
                        total = 0
                        for j, item in enumerate(data_summary, 1):
                            row = table.rows[j]
                            row.cells[0].text = str(j)
                            row.cells[1].text = item['label']
                            value = item['value']
                            if isinstance(value, (int, float)):
                                total += value
                                row.cells[2].text = f"{int(value):,}"
                            else:
                                row.cells[2].text = str(value)
                        
                        # Total row
                        total_row = table.rows[-1]
                        total_row.cells[0].text = ''
                        total_row.cells[1].text = 'TOTAL'
                        total_row.cells[1].paragraphs[0].runs[0].bold = True
                        total_row.cells[2].text = f"{int(total):,}"
                        total_row.cells[2].paragraphs[0].runs[0].bold = True
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"Tipe: {viz.get('type', 'chart')}").italic = True
                    
                    doc.add_paragraph()
            
            # === INSIGHTS ===
            if session_data['insights']:
                doc.add_page_break()
                doc.add_heading('💡 Insight Analisis', level=1)
                
                for i, insight in enumerate(session_data['insights'], 1):
                    insight_text = insight if isinstance(insight, str) else str(insight)
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(insight_text)
            
            # === POLICY RECOMMENDATIONS ===
            if session_data['policies']:
                doc.add_page_break()
                doc.add_heading('🎯 Rekomendasi Kebijakan', level=1)
                
                for i, policy in enumerate(session_data['policies'], 1):
                    # Policy title with priority
                    priority = policy.get('priority', 'medium')
                    priority_icons = {
                        'high': '🔴 Prioritas Tinggi',
                        'medium': '🟡 Prioritas Menengah',
                        'low': '🟢 Prioritas Rendah'
                    }
                    
                    title_text = f"{i}. {policy.get('title', 'Rekomendasi')}"
                    h = doc.add_heading(title_text, level=2)
                    
                    # Priority badge
                    priority_p = doc.add_paragraph()
                    priority_run = priority_p.add_run(priority_icons.get(priority, '🟡 Prioritas Menengah'))
                    priority_run.font.size = Pt(10)
                    priority_run.italic = True
                    
                    # Description
                    description = policy.get('description', '')
                    if description:
                        doc.add_paragraph(description)
                    
                    # Category and Impact
                    category = policy.get('category', '')
                    impact = policy.get('impact', '')
                    
                    if category:
                        p = doc.add_paragraph()
                        p.add_run('Kategori: ').bold = True
                        p.add_run(str(category))
                    
                    if impact:
                        p = doc.add_paragraph()
                        p.add_run('Dampak: ').bold = True
                        p.add_run(impact)
                    
                    # Implementation steps
                    steps = policy.get('implementation_steps', [])
                    if steps:
                        p = doc.add_paragraph()
                        p.add_run('Langkah Implementasi:').bold = True
                        
                        for j, step in enumerate(steps, 1):
                            step_p = doc.add_paragraph(style='List Number')
                            step_p.add_run(step)
                            step_p.paragraph_format.left_indent = Inches(0.5)
                    
                    doc.add_paragraph()
            
            # === FOOTER ===
            doc.add_paragraph()
            doc.add_paragraph('─' * 50)
            
            footer = doc.add_paragraph()
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            footer_run = footer.add_run(
                f'Laporan ini dihasilkan secara otomatis oleh Sistem Analisis Sensus Ekonomi Indonesia.\n'
                f'Data bersumber dari Sensus Ekonomi 2016, Badan Pusat Statistik (BPS).\n'
                f'© {datetime.now().year} - Smart SE26 Agentic AI Chatbot'
            )
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            footer_run.italic = True
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}", exc_info=True)
            raise
    
    def generate_html_report(self, session) -> str:
        """Generate HTML report with embedded chart configurations for web viewing"""
        try:
            session_data = self._extract_session_data(session)
            date_str = datetime.now().strftime("%d %B %Y, %H:%M WIB")
            
            # Build visualizations HTML with ECharts
            viz_html = ""
            chart_scripts = ""
            
            if session_data['visualizations']:
                viz_html = "<section class='section'><h2>📈 Data Visualisasi</h2>"
                
                for i, viz in enumerate(session_data['visualizations']):
                    chart_id = f"chart_{i}"
                    viz_title = viz.get('title', f'Visualisasi {i+1}')
                    config = viz.get('config', {})
                    
                    viz_html += f"""
                    <div class='chart-container'>
                        <h3>{viz_title}</h3>
                        <div id='{chart_id}' class='chart'></div>
                    </div>
                    """
                    
                    config_json = json.dumps(config, ensure_ascii=False)
                    chart_scripts += f"""
                    var {chart_id} = echarts.init(document.getElementById('{chart_id}'));
                    {chart_id}.setOption({config_json});
                    """
                
                viz_html += "</section>"
            
            # Build insights HTML
            insights_html = ""
            if session_data['insights']:
                insights_html = "<section class='section'><h2>💡 Insight Analisis</h2><ul class='insights-list'>"
                for insight in session_data['insights']:
                    insight_text = insight if isinstance(insight, str) else str(insight)
                    insights_html += f"<li>{insight_text}</li>"
                insights_html += "</ul></section>"
            
            # Build policies HTML
            policies_html = ""
            if session_data['policies']:
                policies_html = "<section class='section'><h2>🎯 Rekomendasi Kebijakan</h2>"
                
                for i, policy in enumerate(session_data['policies'], 1):
                    priority = policy.get('priority', 'medium')
                    priority_class = f"priority-{priority}"
                    priority_badges = {
                        'high': '🔴 Prioritas Tinggi',
                        'medium': '🟡 Prioritas Menengah',
                        'low': '🟢 Prioritas Rendah'
                    }
                    
                    steps_html = ""
                    steps = policy.get('implementation_steps', [])
                    if steps:
                        steps_html = "<div class='steps'><strong>Langkah Implementasi:</strong><ol>"
                        for step in steps:
                            steps_html += f"<li>{step}</li>"
                        steps_html += "</ol></div>"
                    
                    policies_html += f"""
                    <div class='policy-card {priority_class}'>
                        <div class='policy-header'>
                            <h3>{i}. {policy.get('title', 'Rekomendasi')}</h3>
                            <span class='priority-badge'>{priority_badges.get(priority, '🟡')}</span>
                        </div>
                        <p class='description'>{policy.get('description', '')}</p>
                        <p class='meta'><strong>Kategori:</strong> {policy.get('category', '-')} | <strong>Dampak:</strong> {policy.get('impact', '-')}</p>
                        {steps_html}
                    </div>
                    """
                
                policies_html += "</section>"
            
            # Build messages HTML
            messages_html = ""
            for msg in session_data['messages']:
                sender_class = 'user' if msg['sender'] == 'user' else 'ai'
                sender_label = '👤 Pengguna' if msg['sender'] == 'user' else '🤖 AI Asisten'
                content = msg['content'].replace('\n', '<br>')
                
                messages_html += f"""
                <div class='message {sender_class}'>
                    <div class='sender'>{sender_label}</div>
                    <div class='content'>{content}</div>
                </div>
                """
            
            # Complete HTML
            html = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Laporan Analisis Sensus Ekonomi</title>
    <script src="https://cdn.jsdelivr.net/npm/echarts@5.4.3/dist/echarts.min.js"></script>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Segoe UI', Tahoma, sans-serif; line-height: 1.6; color: #2c3e50; background: #f8f9fa; }}
        .container {{ max-width: 1000px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #e74c3c, #c0392b); color: white; padding: 30px; text-align: center; border-radius: 10px; margin-bottom: 20px; }}
        .header h1 {{ font-size: 1.8rem; margin-bottom: 10px; }}
        .meta {{ font-size: 0.9rem; opacity: 0.9; }}
        .section {{ background: white; padding: 25px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.05); }}
        .section h2 {{ color: #e74c3c; margin-bottom: 15px; font-size: 1.3rem; }}
        .message {{ padding: 15px; margin-bottom: 15px; border-radius: 8px; }}
        .message.user {{ background: #e8f4f8; border-left: 4px solid #3498db; }}
        .message.ai {{ background: #fef5e7; border-left: 4px solid #f39c12; }}
        .sender {{ font-weight: bold; margin-bottom: 8px; }}
        .chart-container {{ margin-bottom: 25px; padding: 15px; border: 1px solid #eee; border-radius: 8px; }}
        .chart-container h3 {{ color: #2c3e50; margin-bottom: 10px; font-size: 1rem; }}
        .chart {{ width: 100%; height: 350px; }}
        .insights-list {{ list-style: none; }}
        .insights-list li {{ padding: 10px 15px; margin-bottom: 8px; background: #e8f6f3; border-left: 4px solid #1abc9c; border-radius: 4px; }}
        .policy-card {{ padding: 20px; margin-bottom: 15px; border-radius: 8px; background: #fafafa; }}
        .policy-card.priority-high {{ border-left: 4px solid #e74c3c; }}
        .policy-card.priority-medium {{ border-left: 4px solid #f39c12; }}
        .policy-card.priority-low {{ border-left: 4px solid #2ecc71; }}
        .policy-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; }}
        .policy-header h3 {{ font-size: 1.1rem; color: #2c3e50; }}
        .priority-badge {{ font-size: 0.85rem; }}
        .description {{ margin-bottom: 10px; }}
        .meta {{ font-size: 0.9rem; color: #666; margin-bottom: 10px; }}
        .steps {{ background: white; padding: 10px 15px; border-radius: 5px; margin-top: 10px; }}
        .steps ol {{ margin-left: 20px; margin-top: 8px; }}
        .footer {{ text-align: center; padding: 20px; color: #666; font-size: 0.85rem; }}
        @media print {{ .chart {{ page-break-inside: avoid; }} }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📊 Laporan Analisis Sensus Ekonomi Indonesia</h1>
            <p class="meta">Tanggal: {date_str} | Topik: {session_data['title']}</p>
        </div>
        
        <section class="section">
            <h2>📝 Ringkasan Analisis</h2>
            {messages_html}
        </section>
        
        {viz_html}
        {insights_html}
        {policies_html}
        
        <div class="footer">
            <p>Laporan ini dihasilkan secara otomatis oleh Sistem Analisis Sensus Ekonomi Indonesia.<br>
            Data bersumber dari Sensus Ekonomi 2016, BPS. © {datetime.now().year}</p>
        </div>
    </div>
    
    <script>
        {chart_scripts}
        window.addEventListener('resize', function() {{
            var charts = document.querySelectorAll('.chart');
            charts.forEach(function(el) {{
                var chart = echarts.getInstanceByDom(el);
                if (chart) chart.resize();
            }});
        }});
    </script>
</body>
</html>"""
            
            return html
            
        except Exception as e:
            logger.error(f"Error generating HTML report: {e}", exc_info=True)
            raise